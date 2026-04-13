import base64
import io
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from typing import Optional
from services.pdf_parser import extract_pdf_text
from services.docx_parser import extract_docx_text
from services.csv_parser import parse_csv
from services.ai_service import revise_media_plan
from models.schemas import MediaPlanResponse, ChannelPlan, PlanChange, ROIScenario

router = APIRouter()


@router.post("/revise-plan", response_model=MediaPlanResponse)
async def revise_plan(
    media_plan: UploadFile = File(...),
    new_objectives:  str = Form(...),
    new_budget: float = Form(...),
    currency: str = Form("EUR"),
    timeframe: str = Form("1 month"),
    performance_csv: Optional[UploadFile] = File(None),
):
    filename = (media_plan.filename or "").lower()
    file_bytes = await media_plan.read()
    if len(file_bytes) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Media plan file too large. Max 15MB.")

    if filename.endswith(".pdf"):
        plan_text = extract_pdf_text(file_bytes)
    elif filename.endswith(".docx"):
        plan_text = extract_docx_text(file_bytes)
    else:
        raise HTTPException(status_code=400, detail="Media plan must be a PDF or DOCX file.")

    # Optional CSV
    perf_data = None
    if performance_csv and performance_csv.filename:
        csv_bytes = await performance_csv.read()
        if len(csv_bytes) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Performance CSV too large. Max 10MB.")
        parsed = parse_csv(csv_bytes)
        perf_data = parsed["metrics_json"]

    # AI revision
    result = await revise_media_plan(
        original_plan_text=plan_text,
        performance_data=perf_data,
        new_objectives=new_objectives,
        new_budget=new_budget,
        currency=currency,
        timeframe=timeframe,
    )

    # Generate downloadable .docx
    docx_b64 = _generate_docx(result)

    return MediaPlanResponse(
        original_summary=result.get("original_summary", ""),
        revised_channels=[ChannelPlan(**c) for c in result.get("revised_channels", [])],
        total_budget=new_budget,
        changes=[PlanChange(**c) for c in result.get("changes", [])],
        roi_projections=[ROIScenario(**r) for r in result.get("roi_projections", [])],
        executive_summary=result.get("executive_summary", ""),
        docx_base64=docx_b64,
    )


def _generate_docx(plan: dict) -> str:
    """Generates a Word document from the revised plan and returns base64."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading("Revised Media Plan", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(plan.get("executive_summary", ""))

        # Channels
        doc.add_heading("Channel Breakdown", level=1)
        for ch in plan.get("revised_channels", []):
            doc.add_heading(ch["channel"], level=2)
            doc.add_paragraph(f"Budget: {ch['budget']} ({ch['budget_pct']}%)")
            doc.add_paragraph(f"Objective: {ch['objective']}")
            doc.add_paragraph(f"Target Audience: {ch['target_audience']}")
            doc.add_paragraph(f"KPI: {ch['kpi']}")
            doc.add_paragraph(f"Rationale: {ch['rationale']}")

        # Changes
        doc.add_heading("Key Changes", level=1)
        for change in plan.get("changes", []):
            doc.add_heading(change["area"], level=2)
            doc.add_paragraph(f"Original: {change['original']}")
            doc.add_paragraph(f"Revised: {change['revised']}")
            doc.add_paragraph(f"Reason: {change['reason']}")
            doc.add_paragraph(f"Expected Impact: {change['expected_impact']}")

        # ROI Projections
        doc.add_heading("ROI Projections", level=1)
        for proj in plan.get("roi_projections", []):
            doc.add_heading(proj["name"].capitalize(), level=2)
            doc.add_paragraph(f"Projected ROAS: {proj['projected_roas']}x")
            doc.add_paragraph(f"Projected Conversions: {proj['projected_conversions']}")
            doc.add_paragraph(f"Projected Revenue: {proj['projected_revenue']}")
            doc.add_paragraph(f"Assumptions: {proj['assumptions']}")

        buf = io.BytesIO()
        doc.save(buf)
        return base64.b64encode(buf.getvalue()).decode()

    except ImportError:
        return ""
    except Exception:
        return ""
